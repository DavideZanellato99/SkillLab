"""Il testo dentro un documento caricato, qualunque formato abbia.

Un solo posto che sa aprire i formati, perché il resto della simulazione
tecnica lavora su una stringa e non deve sapere se quella stringa veniva da
un PDF o da un file di testo. Quello che esce di qui è normalizzato: niente
righe vuote a ripetizione, niente spazi doppi, niente ritorni carrello di
Windows, così i passaggi in cui il documento viene spezzato hanno la stessa
forma qualunque sia la provenienza.

Il file originale non viene conservato da nessuna parte: dopo questo passaggio
esiste solo il testo (vedi ``TechnicalSimulation.document_text``). Un PDF di
procedura ricaricato è un PDF che si ricarica, non un archivio da tenere.

I due formati binari hanno una libreria ciascuna, e nessuna delle due sa
leggere quello che non è: un PDF fatto di pagine scansionate non contiene
testo, contiene immagini di testo, e da qui esce vuoto. Il chiamante deve
trattare la stringa vuota come un errore da mostrare a chi ha caricato,
perché lì il problema è il file e nessun ritentativo lo cambia.
"""

import io
import re
import zipfile

# Estensioni accettate, con il tipo di lettura che ognuna richiede. Il
# formato si decide da qui e non dal content-type dichiarato dal browser,
# che su Windows arriva vuoto o sbagliato più spesso di quanto si creda.
PDF_EXTENSIONS = (".pdf",)
DOCX_EXTENSIONS = (".docx",)
TEXT_EXTENSIONS = (".txt", ".md", ".markdown")
SUPPORTED_EXTENSIONS = PDF_EXTENSIONS + DOCX_EXTENSIONS + TEXT_EXTENSIONS

# Quanto può pesare un documento caricato. Dieci MB sono una procedura molto
# lunga o un manuale intero: oltre non c'è un caso d'uso, c'è un file
# sbagliato, e il tetto serve a scoprirlo prima di spendere il tempo di
# estrarlo.
MAX_DOCUMENT_BYTES = 10 * 1024 * 1024

# Quanto può diventare grande quel file una volta aperto, e su quante pagine
# può stendersi. Sono il secondo tetto, e serve perché il primo misura la
# cosa sbagliata: un .docx è un archivio compresso, quindi dieci MB di file
# sono centinaia di volte tanto una volta srotolati in memoria, e un PDF di
# poche centinaia di kB può dichiarare decine di migliaia di pagine. In
# entrambi i casi il file passa il controllo sui byte e a cadere è il
# processo che prova a leggerlo, cioè il backend che in quel momento sta
# servendo anche tutti gli altri.
#
# I valori sono larghi di proposito: 200 MB srotolati e 500 pagine stanno
# molto sopra qualunque procedura vera, quindi quello che fermano non è un
# documento lungo, è un documento costruito apposta.
MAX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024
MAX_PDF_PAGES = 500


def _normalize(text: str) -> str:
    """Spazi, ritorni a capo e righe vuote ridotti a una forma sola."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Spazi e tabulazioni ripetuti dentro la riga, che nei PDF sono la norma
    # Fra questi lo spazio unificatore, che PDF e Word seminano ovunque e che
    # a occhio e indistinguibile da uno spazio normale: lasciato passare,
    # spezza le parole in due in ogni confronto e in ogni ricerca.
    text = re.sub(r"[ \t\xa0]+", " ", text)
    # Al massimo una riga vuota di separazione fra due paragrafi
    text = re.sub(r"\n\s*\n\s*(\n\s*)+", "\n\n", text)
    return "\n".join(line.strip() for line in text.split("\n")).strip()


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ValueError(f"Il documento ha troppe pagine: il massimo è {MAX_PDF_PAGES}.")
    # Una riga vuota fra una pagina e l'altra: senza, l'ultima frase di una
    # pagina e il titolo della successiva finiscono nello stesso paragrafo,
    # e da lì nello stesso passaggio.
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _check_zip_bomb(data: bytes) -> None:
    """Rifiuta un archivio che aperto non ci starebbe in memoria.

    Un .docx è uno zip, e uno zip dichiara nel proprio indice quanto pesa
    ogni pezzo una volta srotolato: si legge di lì, prima di srotolare
    davvero, che è tutto il punto. La somma è la misura giusta perché il
    lettore di documenti apre l'XML per intero, quindi quello che conta non
    è il pezzo più grande ma quanto fa il totale.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archivio:
            srotolato = sum(voce.file_size for voce in archivio.infolist())
    except zipfile.BadZipFile:
        raise ValueError("Il documento non è leggibile: il file sembra danneggiato.")
    if srotolato > MAX_UNCOMPRESSED_BYTES:
        raise ValueError(
            "Il documento è troppo grande una volta aperto: "
            f"il massimo è {MAX_UNCOMPRESSED_BYTES // (1024 * 1024)} MB."
        )


def _extract_docx(data: bytes) -> str:
    from docx import Document

    _check_zip_bomb(data)
    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # Le tabelle non stanno fra i paragrafi e sparirebbero in silenzio, che
    # in una procedura è proprio dove stanno i passaggi operativi.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_plain(data: bytes) -> str:
    # UTF-8 è quello che scrivono tutti; latin-1 come rete di sicurezza per i
    # file salvati da un editor di Windows, e accetta qualunque byte, quindi
    # da qui non esce mai un errore di decodifica.
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1")


def is_supported(filename: str) -> bool:
    return filename.lower().endswith(SUPPORTED_EXTENSIONS)


def extract_text(filename: str, data: bytes) -> str:
    """Il testo del documento, normalizzato. Stringa vuota se non ce n'è.

    Solleva ValueError se l'estensione non è fra quelle supportate, e lascia
    passare gli errori delle librerie di lettura: un file corrotto o protetto
    da password è un errore da riferire a chi lo ha caricato, non qualcosa da
    trasformare in un documento vuoto.
    """
    name = filename.lower()
    if name.endswith(PDF_EXTENSIONS):
        raw = _extract_pdf(data)
    elif name.endswith(DOCX_EXTENSIONS):
        raw = _extract_docx(data)
    elif name.endswith(TEXT_EXTENSIONS):
        raw = _extract_plain(data)
    else:
        raise ValueError(f"Formato non supportato: {filename}")
    return _normalize(raw)
