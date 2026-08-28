"""Il testo che si tira fuori da un documento caricato.

I due formati binari hanno una libreria ciascuna, e da qui in poi la
simulazione tecnica lavora su una stringa: se l'estrazione perde un pezzo,
le domande nascono su un documento che non è quello caricato, e il super
admin che le rilegge non ha modo di accorgersene, perché una domanda su
metà procedura sembra comunque una domanda.

I documenti di prova si costruiscono qui invece di tenerli come file
allegati: un .docx nel repository è uno zip che nessuno può rileggere in una
differenza, e quello che si vuole verificare, che le tabelle non spariscano,
si legge molto meglio nelle righe che lo costruiscono.
"""

import io

import pytest
from docx import Document
from fpdf import FPDF

import document_text
from document_text import extract_text


def _pdf(*pagine: str) -> bytes:
    documento = FPDF()
    documento.set_font("helvetica", size=12)
    for testo in pagine:
        documento.add_page()
        documento.multi_cell(0, 10, testo)
    return bytes(documento.output())


def _docx(paragrafi: list[str], tabelle: list[list[list[str]]] | None = None) -> bytes:
    documento = Document()
    for testo in paragrafi:
        documento.add_paragraph(testo)
    for righe in tabelle or []:
        tabella = documento.add_table(rows=len(righe), cols=len(righe[0]))
        for riga, valori in zip(tabella.rows, righe, strict=True):
            for cella, valore in zip(riga.cells, valori, strict=True):
                cella.text = valore
    buffer = io.BytesIO()
    documento.save(buffer)
    return buffer.getvalue()


# ── PDF ───────────────────────────────────────────────────────────────


def test_il_pdf_esce_con_tutte_le_sue_pagine():
    testo = extract_text("procedura.pdf", _pdf("Prima pagina.", "Seconda pagina."))

    assert "Prima pagina." in testo
    assert "Seconda pagina." in testo


def test_fra_una_pagina_e_l_altra_resta_uno_stacco():
    """Senza, l'ultima frase di una pagina e il titolo della successiva
    finiscono nello stesso paragrafo, e da lì nello stesso passaggio."""
    testo = extract_text("procedura.pdf", _pdf("Fine del capitolo.", "Titolo nuovo."))

    assert "Fine del capitolo.\n\nTitolo nuovo." in testo


def test_un_pdf_di_pagine_scansionate_esce_vuoto():
    """Un PDF fatto di immagini non contiene testo, e da qui esce la stringa
    vuota: è il chiamante a doverla trattare come un file da rifiutare, e
    non c'è ritentativo che la cambi."""
    documento = FPDF()
    documento.add_page()

    assert extract_text("scansione.pdf", bytes(documento.output())) == ""


def test_un_pdf_illeggibile_risale_invece_di_diventare_un_documento_vuoto():
    """Un file corrotto o protetto da password è un errore da riferire a chi
    lo ha caricato: se uscisse vuoto, la simulazione si genererebbe su
    niente."""
    with pytest.raises(Exception):  # noqa: B017 (è pypdf a scegliere il tipo)
        extract_text("rotto.pdf", b"non sono un pdf")


# ── DOCX ──────────────────────────────────────────────────────────────


def test_il_docx_esce_con_i_suoi_paragrafi():
    testo = extract_text("manuale.docx", _docx(["Primo paragrafo.", "Secondo paragrafo."]))

    assert testo == "Primo paragrafo.\nSecondo paragrafo."


def test_le_tabelle_del_docx_non_spariscono():
    """In una procedura aziendale è proprio nelle tabelle che stanno i
    passaggi operativi: perderle in silenzio lascerebbe un documento
    apparentemente completo senza la parte che conta."""
    testo = extract_text(
        "manuale.docx",
        _docx(
            ["Casistiche e uffici."],
            tabelle=[[["Reclamo", "Ufficio reclami"], ["Rimborso", "Amministrazione"]]],
        ),
    )

    assert "Reclamo | Ufficio reclami" in testo
    assert "Rimborso | Amministrazione" in testo


def test_le_celle_vuote_non_lasciano_separatori_a_vuoto():
    testo = extract_text("manuale.docx", _docx([], tabelle=[[["Reclamo", ""]]]))

    assert "Reclamo" in testo
    assert "|" not in testo


# ── Testo semplice ────────────────────────────────────────────────────


def test_un_file_salvato_da_un_editor_di_windows_si_legge_lo_stesso():
    """Latin-1 accetta qualunque byte, quindi da qui non esce mai un errore
    di decodifica: un accento sbagliato è meglio di un caricamento
    rifiutato."""
    testo = extract_text("note.txt", "Perché però".encode("latin-1"))

    assert "Perch" in testo


def test_il_markdown_e_un_documento_come_gli_altri():
    assert extract_text("procedura.md", b"# Titolo\n\nContenuto.") == "# Titolo\n\nContenuto."


def test_un_formato_che_non_si_sa_aprire_si_rifiuta_subito():
    with pytest.raises(ValueError, match="Formato non supportato"):
        extract_text("foglio.xlsx", b"qualunque cosa")


def test_l_estensione_decide_il_formato_a_prescindere_da_come_e_scritta():
    """Il tipo dichiarato dal browser su Windows arriva vuoto o sbagliato più
    spesso di quanto si creda, quindi qui si guarda solo il nome."""
    assert extract_text("PROCEDURA.TXT", b"Contenuto.") == "Contenuto."


# ── I file costruiti per far cadere chi li apre ───────────────────────


def test_un_docx_che_srotolato_non_ci_starebbe_si_rifiuta(monkeypatch):
    """Il tetto sui byte del file misura la cosa sbagliata: un .docx è un
    archivio compresso, quindi dieci MB di file sono centinaia di volte
    tanto una volta aperti. La misura giusta la dichiara l'archivio stesso
    nel proprio indice, e si legge di lì prima di srotolare davvero."""
    import zipfile

    monkeypatch.setattr(document_text, "MAX_UNCOMPRESSED_BYTES", 1024)
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archivio:
        # Comprime benissimo: pochi byte sul disco, molti una volta aperti.
        archivio.writestr("word/document.xml", b"a" * 100_000)

    with pytest.raises(ValueError, match="troppo grande una volta aperto"):
        extract_text("bomba.docx", buffer.getvalue())


def test_un_docx_danneggiato_lo_dice_invece_di_scoppiare():
    with pytest.raises(ValueError, match="danneggiato"):
        extract_text("rotto.docx", b"questo non e uno zip")


def test_un_pdf_con_troppe_pagine_si_rifiuta(monkeypatch):
    """Un PDF di poche centinaia di kB può dichiararne decine di migliaia, e
    a cadere non è il file, è il processo che prova a leggerlo."""
    monkeypatch.setattr(document_text, "MAX_PDF_PAGES", 2)
    pdf = FPDF()
    for numero in range(3):
        pdf.add_page()
        pdf.set_font("helvetica", size=12)
        pdf.cell(0, 10, f"Pagina {numero}")

    with pytest.raises(ValueError, match="troppe pagine"):
        extract_text("lungo.pdf", bytes(pdf.output()))
